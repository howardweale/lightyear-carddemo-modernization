#!/usr/bin/env python3
"""Build and verify the MS #1-47 customer documentation library."""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import shutil
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = ROOT / "docs" / "milestones"
CATALOG_PATH = DOC_ROOT / "catalog.json"
CHANGELOG_PATH = ROOT / "CHANGELOG.md"
MANIFEST_PATH = DOC_ROOT / "manifest.json"
BRAND_ROOT = ROOT / "brand"
BRAND_ASSETS = BRAND_ROOT / "assets"
BRAND_LOGO_SVG = BRAND_ASSETS / "lightyear-primary.svg"
BRAND_LOGO_PNG = BRAND_ASSETS / "lightyear-primary.png"
GENERATOR_VERSION = "1.3"
REPOSITORY = "howardweale/lightyear-carddemo-modernization"
DEFAULT_BRANCH = "main"
GITHUB_BLOB_ROOT = f"https://github.com/{REPOSITORY}/blob/{DEFAULT_BRANCH}"
GITHUB_RAW_ROOT = f"https://raw.githubusercontent.com/{REPOSITORY}/{DEFAULT_BRANCH}"
PAGES_INDEX = f"https://howardweale.github.io/{REPOSITORY.split('/', 1)[1]}/milestones/"
FIXED_TIME = datetime(2026, 8, 31, 12, 0, 0, tzinfo=timezone.utc)
EXPECTED_MILESTONES = tuple(range(1, 48))
EXPECTED_ARTIFACTS = len(EXPECTED_MILESTONES) * 3
BOUNDARY_TERMS = (
    "remain false", "remains false", "remain blocked", "remains blocked",
    "unclaimed", "not claim", "does not", "no customer", "non-production",
    "kept native", "kept all live", "kept every", "cannot satisfy", "excluded",
)


def sha256_path(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def github_blob(path: str) -> str:
    return f"{GITHUB_BLOB_ROOT}/{path}"


def github_raw(path: str) -> str:
    return f"{GITHUB_RAW_ROOT}/{path}"


def load_catalog() -> dict[str, Any]:
    catalog = json.loads(CATALOG_PATH.read_text(encoding="utf-8"))
    numbers = [entry["number"] for entry in catalog["milestones"]]
    if numbers != list(EXPECTED_MILESTONES):
        raise ValueError(f"catalog must contain MS #1-{EXPECTED_MILESTONES[-1]} exactly; found {numbers}")
    return catalog


def parse_changelog() -> dict[int, list[dict[str, Any]]]:
    sections: dict[int, list[dict[str, Any]]] = {}
    current: dict[str, Any] | None = None
    current_item: list[str] | None = None
    header = re.compile(r"^## 0\.(\d+)\.(\d+)\s+.+?\s+(\d{4}-\d{2}-\d{2})$")
    for raw in CHANGELOG_PATH.read_text(encoding="utf-8").splitlines():
        match = header.match(raw)
        if match:
            minor, patch, date = int(match.group(1)), int(match.group(2)), match.group(3)
            current = {"version": f"0.{minor}.{patch}", "patch": patch, "date": date, "items": []}
            sections.setdefault(minor, []).append(current)
            current_item = None
        elif current is not None and raw.startswith("- "):
            current_item = [raw[2:].strip()]
            current["items"].append(current_item)
        elif current_item is not None and raw.startswith("  "):
            current_item.append(raw.strip())
        elif not raw.strip():
            current_item = None
    normalized: dict[int, list[dict[str, Any]]] = {}
    for minor, releases in sections.items():
        normalized[minor] = []
        for release in sorted(releases, key=lambda row: row["patch"]):
            normalized[minor].append({**release, "items": [" ".join(parts) for parts in release["items"]]})
    return normalized


def plain(text: str) -> str:
    return (
        re.sub(r"`([^`]+)`", r"\1", text)
        .replace("\u2192", "to")
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2011", "-")
    )


def finish_sentence(text: str) -> str:
    value = plain(text).strip()
    return value if value.endswith((".", "!", "?")) else value + "."


def build_model(
    entry: dict[str, Any], releases: dict[int, list[dict[str, Any]]], titles: dict[int, str]
) -> dict[str, Any]:
    number = entry["number"]
    release_rows = releases.get(number, [])
    deliverables = list(entry.get("deliverables", [])) or [
        item for release in release_rows for item in release["items"]
    ]
    if not deliverables:
        raise ValueError(f"MS #{number} has no deliverables")
    if "release" in entry:
        release_label, date_label = entry["release"], entry["date"]
    else:
        release_label = ", ".join(row["version"] for row in release_rows)
        date_label = (
            f"{release_rows[0]['date']} to {release_rows[-1]['date']}"
            if len(release_rows) > 1 else release_rows[0]["date"]
        )
    boundaries = list(entry.get("boundaries", [])) or [
        item for item in deliverables if any(term in item.lower() for term in BOUNDARY_TERMS)
    ]
    if not boundaries:
        boundaries = [
            "This milestone establishes only the bounded capabilities and evidence listed here; later capabilities are not retroactive.",
            "Production readiness and platform equivalence require the explicit evidence and policy gates applicable to the customer environment.",
        ]
    relationship = entry.get("relationship") or (
        "This is the starting engineering milestone. It creates the executable behavioral reference that later graph, factory, qualification, and customer-pilot work can govern."
        if number == 1
        else f"MS #{number} builds on MS #{number - 1} ({titles[number - 1]}). The earlier milestone established its own bounded capability; this milestone adds {entry['title'].lower()} while preserving the earlier evidence and its limits."
    )
    return {
        **entry,
        "status": "Complete",
        "release": release_label,
        "date": date_label,
        "deliverables": [finish_sentence(item) for item in deliverables],
        "boundaries": [finish_sentence(item) for item in boundaries],
        "relationship": relationship,
        "executive_summary": (
            f"MS #{number} - {entry['title']} converts a specific part of the LIGHTYEAR modernization approach "
            f"into a governed, reviewable capability. {entry['customer_value']}"
        ),
    }


def markdown_text(model: dict[str, Any], audience: str) -> str:
    number = model["number"]
    lines = [
        "![LIGHTYEAR primary logo](../../../brand/assets/lightyear-primary.svg)", "",
        "*Where context becomes trusted action.*", "",
        f"# MS #{number:02d} - {model['title']}", "",
        f"> **Status:** {model['status']}", ">", f"> **Release:** {model['release']}", ">",
        f"> **Date:** {model['date']}", ">", f"> **Audience:** {audience}", "",
        "## Executive summary", "", model["executive_summary"], "",
        "## Why this matters to a customer", "", model["customer_value"], "",
        "The practical result is a narrower, evidence-backed decision instead of a broad modernization claim. "
        "The milestone packages capability, proof, and limitations together so commercial, architecture, "
        "delivery, security, and assurance teams can review the same record.", "",
        "## What the milestone delivers", "",
    ]
    lines.extend(f"- {item}" for item in model["deliverables"])
    lines.extend([
        "", "## How it differs from earlier work", "", model["relationship"], "",
        "Earlier work remains useful evidence, but it is not automatically promoted into this milestone. "
        "A successful parser, simulator, local comparison, or generated artifact is not presented as live "
        "platform equivalence or production readiness.", "",
        "## What a customer can do with it", "",
        f"A customer can use {model['title'].lower()} as a bounded decision and delivery capability. "
        "It can be attached to the customer's approved source package, owners, policies, target architecture, "
        "acceptance criteria, and authorized evidence.", "",
        "Unresolved semantic loss, policy decisions, unsupported behavior, and live-evidence requirements remain "
        "visible rather than being averaged into one pass/fail statement.", "",
        "## Evidence and acceptance posture", "",
        "The milestone is complete as a repository capability because its committed implementation and release "
        "record are present. Customer qualification remains claim-specific: only the delivered behaviors listed "
        "above are supported, and only at the evidence class recorded by their underlying receipts and gates.", "",
        "A customer review should confirm:", "",
        "1. the exact source, graph, policy, fixture, and artifact identities referenced by the evidence;",
        "2. the difference between deterministic development evidence and authorized live evidence;",
        "3. every policy decision, lossy transformation, unsupported behavior, and explicit exclusion;",
        "4. the acceptance gates that passed and the live or production gates that remain blocked; and",
        "5. that later changes have not invalidated the content-addressed evidence.", "",
        "## Boundaries and claims not made", "",
    ])
    lines.extend(f"- {item}" for item in model["boundaries"])
    lines.extend([
        "", "These boundaries are part of the deliverable. They prevent bounded development evidence from becoming "
        "an unsupported customer promise.", "", "## Source of record", "",
        f"- Release record: [CHANGELOG.md]({github_blob('CHANGELOG.md')}), release section(s) {model['release']}.",
        f"- Product and operating guidance: [README.md]({github_blob('README.md')}).",
        f"- Governing sequence and claim classifications: [LIGHTYEAR-ROADMAP.md]({github_blob('LIGHTYEAR-ROADMAP.md')}).",
        f"- Machine-readable milestone metadata: [catalog.json]({github_blob('docs/milestones/catalog.json')}).",
        f"- Artifact integrity: [manifest.json]({github_blob('docs/milestones/manifest.json')}).", "", "---", "",
        "This document is a customer-readable explanation of committed repository evidence. The underlying schemas, "
        "receipts, ledgers, tests, and policy decisions remain authoritative when greater detail is required.", "",
    ])
    return "\n".join(lines)


def set_run_font(run: Any, name: str = "Arial") -> None:
    from docx.oxml.ns import qn
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)


def set_table_geometry(table: Any, widths: list[int], indent: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    total = sum(widths)
    table.autofit = False
    properties = table._tbl.tblPr
    for name in ("tblW", "tblInd", "tblLayout"):
        for child in list(properties):
            if child.tag == qn(f"w:{name}"):
                properties.remove(child)
    for name, attrs in (
        ("tblW", {"w:w": str(total), "w:type": "dxa"}),
        ("tblInd", {"w:w": str(indent), "w:type": "dxa"}),
        ("tblLayout", {"w:type": "fixed"}),
    ):
        element = OxmlElement(f"w:{name}")
        for key, value in attrs.items():
            element.set(qn(key), value)
        properties.append(element)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                part = margins.find(qn(f"w:{side}"))
                if part is None:
                    part = OxmlElement(f"w:{side}")
                    margins.append(part)
                part.set(qn("w:w"), str(value))
                part.set(qn("w:type"), "dxa")


def shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_page_field(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    for tag, value in (("fldChar", "begin"), ("instrText", " PAGE "), ("fldChar", "separate"), ("t", "1"), ("fldChar", "end")):
        node = OxmlElement(f"w:{tag}")
        if tag == "fldChar":
            node.set(qn("w:fldCharType"), value)
        else:
            node.text = value
        run._r.append(node)


def normalize_docx(source: Path, output: Path) -> None:
    with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
        for name in sorted(zin.namelist()):
            info = zipfile.ZipInfo(name, (2026, 8, 31, 12, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            zout.writestr(info, zin.read(name))


def build_docx(model: dict[str, Any], audience: str, output: Path) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("DOCX build requires python-docx; install the docs optional dependencies") from exc
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = Inches(0.88)
    section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance, section.footer_distance = Inches(0.28), Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Arial", Pt(10.25)
    normal.font.color.rgb = RGBColor.from_string("15184D")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(5), 1.08
    for name, size, color, before, after in (
        ("Heading 1", 15.5, "7D57EA", 13, 6), ("Heading 2", 13, "7D57EA", 11, 5),
        ("Heading 3", 12, "15184D", 8, 4),
    ):
        style = document.styles[name]
        style.font.name, style.font.size, style.font.bold = "Arial", Pt(size), True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name, style.font.size = "Arial", Pt(10.25)
        style.paragraph_format.left_indent, style.paragraph_format.first_line_indent = Inches(0.5), Inches(-0.25)
        style.paragraph_format.space_after, style.paragraph_format.line_spacing = Pt(5), 1.10
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(2)
    logo_run = header.add_run()
    logo_run.add_picture(str(BRAND_LOGO_PNG), width=Inches(1.05))
    run = header.add_run("   MILESTONE DOCUMENTATION")
    set_run_font(run)
    run.font.size, run.font.bold, run.font.color.rgb = Pt(8), True, RGBColor.from_string("676985")
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, value in (("w:val", "single"), ("w:sz", "7"), ("w:space", "3"), ("w:color", "A7702C")):
        bottom.set(qn(key), value)
    border.append(bottom)
    header._p.get_or_add_pPr().append(border)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    run = footer.add_run("Underlying evidence remains authoritative.")
    set_run_font(run)
    run.font.size, run.font.color.rgb = Pt(7.5), RGBColor.from_string("676985")
    run = footer.add_run(f"\tMS #{model['number']:02d}  |  Page ")
    set_run_font(run)
    run.font.size, run.font.color.rgb = Pt(8.5), RGBColor.from_string("676985")
    add_page_field(footer)
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(0)
    run = kicker.add_run("CUSTOMER MILESTONE BRIEF")
    set_run_font(run)
    run.font.size, run.font.bold, run.font.color.rgb = Pt(9), True, RGBColor.from_string("A7702C")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(f"MS #{model['number']:02d} - {plain(model['title'])}")
    set_run_font(run)
    run.font.size, run.font.bold, run.font.color.rgb = Pt(28), True, RGBColor.from_string("15184D")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Purpose, customer value, delivered capability, evidence, and claim boundaries")
    set_run_font(run)
    run.font.size, run.font.color.rgb = Pt(13.5), RGBColor.from_string("676985")
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    metadata = (("Status", model["status"]), ("Release", model["release"]), ("Date", model["date"]), ("Audience", audience))
    for row, (label, value) in zip(table.rows, metadata):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(row.cells[0], "EFEBFB")
        label_run = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(label_run)
        label_run.bold, label_run.font.color.rgb = True, RGBColor.from_string("15184D")
        value_run = row.cells[1].paragraphs[0].add_run(plain(value))
        set_run_font(value_run)
    content = [
        ("Executive summary", [model["executive_summary"]], False),
        ("Why this matters to a customer", [model["customer_value"], "The practical result is a narrower, evidence-backed decision instead of a broad modernization claim. Commercial, architecture, delivery, security, and assurance teams can review the same capability, proof, and limits."], False),
        ("What the milestone delivers", model["deliverables"], True),
        ("How it differs from earlier work", [model["relationship"], "Earlier evidence is not automatically promoted into this milestone. A parser, simulator, local comparison, or generated artifact is not live platform equivalence or production readiness."], False),
        ("What a customer can do with it", [f"A customer can use {model['title'].lower()} as a bounded decision and delivery capability, attached to approved sources, owners, policies, target architecture, acceptance criteria, and authorized evidence.", "Unresolved semantic loss, policy decisions, unsupported behavior, and live-evidence requirements remain visible."], False),
        ("Evidence and acceptance posture", ["The repository capability is complete because its committed implementation and release record are present. Customer qualification remains claim-specific and evidence-class-specific.", "Review exact source and artifact identities, distinguish development from authorized live evidence, inspect policy and loss classifications, confirm passed and blocked gates, and verify that later changes have not invalidated the evidence."], False),
        ("Boundaries and claims not made", model["boundaries"], True),
        ("Source of record", [f"CHANGELOG.md release section(s): {model['release']}", "README.md product and operating guidance", "LIGHTYEAR-ROADMAP.md governing sequence and claim classifications", "docs/milestones/catalog.json metadata and docs/milestones/manifest.json artifact integrity"], True),
    ]
    for heading, values, bullets in content:
        document.add_heading(heading, level=1)
        for value in values:
            document.add_paragraph(plain(value), style="List Bullet" if bullets else None)
    properties = document.core_properties
    properties.title, properties.subject, properties.author = f"MS #{model['number']:02d} - {plain(model['title'])}", "LIGHTYEAR customer milestone documentation", "LIGHTYEAR"
    properties.created = properties.modified = FIXED_TIME.replace(tzinfo=None)
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="milestone-docx-") as temp:
        raw = Path(temp) / output.name
        document.save(raw)
        normalize_docx(raw, output)


def build_pdf(model: dict[str, Any], audience: str, output: Path) -> None:
    try:
        import reportlab
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError("PDF build requires reportlab; install the docs optional dependencies") from exc
    font_dir = Path(reportlab.__file__).resolve().parent / "fonts"
    for name, filename in (("LYSans", "Vera.ttf"), ("LYSans-Bold", "VeraBd.ttf"), ("LYSans-Italic", "VeraIt.ttf")):
        if name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(name, str(font_dir / filename)))
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="LYSans", fontSize=9.7, leading=12.1, spaceAfter=4.5, textColor=colors.HexColor("#15184D"))
    heading = ParagraphStyle("H1", parent=styles["Heading1"], fontName="LYSans-Bold", fontSize=14, leading=16.5, textColor=colors.HexColor("#7D57EA"), spaceBefore=10, spaceAfter=5, keepWithNext=True)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=18, firstLineIndent=-10, spaceAfter=4)
    label = ParagraphStyle("Label", parent=body, fontName="LYSans-Bold", textColor=colors.HexColor("#15184D"), spaceAfter=0)
    value = ParagraphStyle("Value", parent=body, spaceAfter=0)
    def safe(text: str) -> str:
        return html.escape(plain(text))
    def furniture(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.drawImage(str(BRAND_LOGO_PNG), inch, LETTER[1] - 0.66 * inch, width=0.96 * inch, height=0.587 * inch, preserveAspectRatio=True, mask="auto")
        canvas.setStrokeColor(colors.HexColor("#A7702C")); canvas.setLineWidth(0.65)
        canvas.line(inch, LETTER[1] - 0.58 * inch, LETTER[0] - inch, LETTER[1] - 0.58 * inch)
        canvas.setFillColor(colors.HexColor("#676985")); canvas.setFont("LYSans-Bold", 7.5)
        canvas.drawRightString(LETTER[0] - inch, LETTER[1] - 0.45 * inch, "MILESTONE DOCUMENTATION")
        canvas.setFont("LYSans", 8)
        canvas.drawString(inch, 0.48 * inch, "Underlying evidence remains authoritative.")
        canvas.drawRightString(LETTER[0] - inch, 0.48 * inch, f"MS #{model['number']:02d}  |  Page {doc.page}")
        canvas.restoreState()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(output), pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=0.86 * inch, bottomMargin=0.66 * inch, title=f"MS #{model['number']:02d} - {plain(model['title'])}", author="LIGHTYEAR", subject="LIGHTYEAR customer milestone documentation", invariant=1, pageCompression=1)
    story: list[Any] = [Spacer(1, 8)]
    story.append(Paragraph("CUSTOMER MILESTONE BRIEF", ParagraphStyle("Kicker", parent=body, fontName="LYSans-Bold", fontSize=8.5, textColor=colors.HexColor("#A7702C"), spaceAfter=2)))
    story.append(Paragraph(f"MS #{model['number']:02d} - {safe(model['title'])}", ParagraphStyle("Title", parent=styles["Title"], fontName="LYSans-Bold", fontSize=23, leading=27, textColor=colors.HexColor("#15184D"), spaceAfter=6)))
    story.append(Paragraph("Purpose, customer value, delivered capability, evidence, and claim boundaries", ParagraphStyle("Subtitle", parent=body, fontSize=11.5, leading=14, textColor=colors.HexColor("#676985"), spaceAfter=12)))
    rows = [[Paragraph(k, label), Paragraph(safe(v), value)] for k, v in (("Status", model["status"]), ("Release", model["release"]), ("Date", model["date"]), ("Audience", audience))]
    table = Table(rows, colWidths=[1.55 * inch, 4.95 * inch], hAlign="LEFT")
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#DDD7F2")), ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EFEBFB")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    story.append(table)
    content = [
        ("Executive summary", [model["executive_summary"]], False),
        ("Why this matters to a customer", [model["customer_value"], "The practical result is a narrower, evidence-backed decision instead of a broad modernization claim. Commercial, architecture, delivery, security, and assurance teams can review the same capability, proof, and limits."], False),
        ("What the milestone delivers", model["deliverables"], True),
        ("How it differs from earlier work", [model["relationship"], "Earlier evidence is not automatically promoted into this milestone. A parser, simulator, local comparison, or generated artifact is not live platform equivalence or production readiness."], False),
        ("What a customer can do with it", [f"A customer can use {model['title'].lower()} as a bounded decision and delivery capability, attached to approved sources, owners, policies, target architecture, acceptance criteria, and authorized evidence.", "Unresolved semantic loss, policy decisions, unsupported behavior, and live-evidence requirements remain visible."], False),
        ("Evidence and acceptance posture", ["The repository capability is complete because its committed implementation and release record are present. Customer qualification remains claim-specific and evidence-class-specific.", "Review exact source and artifact identities, distinguish development from authorized live evidence, inspect policy and loss classifications, confirm passed and blocked gates, and verify that later changes have not invalidated the evidence."], False),
        ("Boundaries and claims not made", model["boundaries"], True),
        ("Source of record", [f"CHANGELOG.md release section(s): {model['release']}", "README.md product and operating guidance", "LIGHTYEAR-ROADMAP.md governing sequence and claim classifications", "docs/milestones/catalog.json metadata and docs/milestones/manifest.json artifact integrity"], True),
    ]
    for title, values, bullets in content:
        story.append(Paragraph(safe(title), heading))
        for item in values:
            story.append(Paragraph(("-  " if bullets else "") + safe(item), bullet if bullets else body))
    doc.build(story, onFirstPage=furniture, onLaterPages=furniture)


def source_sha256() -> str:
    digest = hashlib.sha256()
    for path in (Path(__file__).resolve(), CATALOG_PATH, CHANGELOG_PATH, BRAND_ROOT / "tokens.json", BRAND_LOGO_SVG, BRAND_LOGO_PNG):
        digest.update(path.relative_to(ROOT).as_posix().encode()); digest.update(b"\0")
        digest.update(path.read_bytes()); digest.update(b"\0")
    return digest.hexdigest()


def format_links(model: dict[str, Any]) -> tuple[str, str, str]:
    stem = f"MS-{model['number']:02d}"
    root = f"docs/milestones/{stem}/{stem}"
    return github_blob(f"{root}.md"), github_raw(f"{root}.docx"), github_blob(f"{root}.pdf")


def write_markdown_index(models: list[dict[str, Any]]) -> None:
    lines = [
        "![LIGHTYEAR primary logo](../../brand/assets/lightyear-primary.svg)", "",
        "*Where context becomes trusted action.*", "",
        "# LIGHTYEAR milestone documentation library", "",
        f"This library is the customer-readable body of record for MS #1 through MS #{EXPECTED_MILESTONES[-1]}. Every milestone is",
        "published from one governed catalog in Markdown, Microsoft Word (`.docx`), and PDF.", "",
        f"**[Open the searchable milestone index]({PAGES_INDEX})** to filter by milestone number, title,",
        "customer value, capability, release, or roadmap phase.", "",
        "The documents explain purpose, customer value, delivered capability, evidence posture, limitations, and",
        "the relationship to earlier work. They do not override the underlying schemas, receipts, ledgers, tests,",
        "or policy decisions.", "", "## Milestone index", "",
        "| Milestone | Title | Status | Formats |", "|---|---|---|---|",
    ]
    for model in models:
        markdown, word, pdf = format_links(model)
        lines.append(f"| MS #{model['number']:02d} | {model['title']} | {model['status']} | [Markdown]({markdown}) - [Download Word]({word}) - [PDF]({pdf}) |")
    lines.extend(["", "## Build and verification", "", "```bash", "./milestone-documentation.sh verify", "./milestone-documentation.sh build", "```", "", "Windows:", "", "```powershell", ".\\milestone-documentation.ps1 verify", ".\\milestone-documentation.ps1 build", "```", "", "`verify` uses only the Python standard library. `build` requires the `docs` optional dependency set.", "", f"The content-addressed [manifest]({github_blob('docs/milestones/manifest.json')}) fails verification if a canonical source changes, an artifact is missing or modified, or an untracked milestone artifact appears.", ""])
    (DOC_ROOT / "README.md").write_text("\n".join(lines), encoding="utf-8")


def roadmap_phase(number: int) -> tuple[str, str]:
    if number <= 10:
        return "foundation", "Foundation"
    if number <= 20:
        return "trust-runtime", "Trust and runtime"
    if number <= 32:
        return "pilot-preparation", "Pilot preparation"
    return "qualification", "Qualification"


def write_html_index(models: list[dict[str, Any]]) -> None:
    rows = []
    for model in models:
        markdown, word, pdf = format_links(model)
        phase_value, phase_label = roadmap_phase(model["number"])
        searchable = " ".join(
            [
                f"MS {model['number']} {model['number']:02d}", model["title"], model["customer_value"],
                model["release"], phase_label, *model["deliverables"], *model["boundaries"],
            ]
        ).lower()
        rows.append(
            f'''<tr class="milestone" data-phase="{phase_value}" data-search="{html.escape(searchable, quote=True)}">
              <td class="number"><span>MS #{model['number']:02d}</span></td>
              <td><a class="title" href="{markdown}">{html.escape(model['title'])}</a><p>{html.escape(model['customer_value'])}</p></td>
              <td class="phase"><span>{phase_label}</span><small>{html.escape(model['release'])}</small></td>
              <td class="formats"><a href="{markdown}">Read</a><a href="{pdf}">PDF</a><a href="{word}" download>Word</a></td>
            </tr>'''
        )
    page = f'''<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <meta name="description" content="Search the governed LIGHTYEAR documentation for milestones MS #1 through MS #{EXPECTED_MILESTONES[-1]}.">
  <title>LIGHTYEAR milestone library</title>
  <style>
    :root {{ color-scheme: light; --ink:#15184d; --muted:#676985; --line:#ddd7f2; --paper:#fefefe; --wash:#f7f6fc; --navy:#15184d; --violet:#7d57ea; --violet-dark:#6942d6; --lavender:#efebfb; --bronze:#a7702c; }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; color:var(--ink); background:var(--wash); font:16px/1.5 Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; }}
    a {{ color:var(--violet-dark); }}
    header {{ background:linear-gradient(135deg, var(--navy), #242866); color:white; padding:3.1rem max(1.25rem, calc((100vw - 1180px)/2)) 3.7rem; }}
    .brand {{ display:flex; align-items:center; gap:1rem; margin-bottom:2rem; }}
    .brand img {{ display:block; width:150px; height:auto; }}
    .brand span {{ padding-left:1rem; color:#d8ccf8; border-left:1px solid rgba(216,204,248,.55); font-size:.92rem; font-weight:700; }}
    header p {{ max-width:760px; margin:.7rem 0 0; color:#ded9f0; font-size:1.08rem; }}
    .eyebrow {{ margin:0; color:#bdaafd; font-size:.78rem; font-weight:800; letter-spacing:.14em; text-transform:uppercase; }}
    h1 {{ margin:.25rem 0 0; font-size:clamp(2rem, 5vw, 3.6rem); line-height:1.05; letter-spacing:-.035em; }}
    main {{ width:min(1180px, calc(100% - 2rem)); margin:-1.35rem auto 3rem; }}
    .controls {{ display:grid; grid-template-columns:minmax(250px, 1fr) minmax(190px, 260px); gap:1rem; padding:1.1rem; background:var(--paper); border:1px solid var(--line); border-radius:14px; box-shadow:0 12px 34px rgba(21,24,77,.1); }}
    label {{ display:block; margin:0 0 .35rem; font-size:.78rem; font-weight:800; letter-spacing:.06em; text-transform:uppercase; }}
    input, select {{ width:100%; min-height:46px; border:1px solid #b8abe8; border-radius:8px; padding:.65rem .8rem; color:var(--ink); background:white; font:inherit; }}
    input:focus, select:focus {{ outline:3px solid rgba(125,87,234,.22); border-color:var(--violet); }}
    .summary {{ display:flex; justify-content:space-between; gap:1rem; align-items:center; margin:1.4rem .2rem .65rem; color:var(--muted); }}
    #result-count {{ font-weight:750; color:var(--ink); }}
    .summary a {{ font-size:.92rem; }}
    .table-wrap {{ overflow:hidden; border:1px solid var(--line); border-radius:12px; background:var(--paper); box-shadow:0 8px 26px rgba(21,24,77,.06); }}
    table {{ width:100%; border-collapse:collapse; }}
    th {{ padding:.78rem 1rem; color:#eee9ff; background:var(--navy); font-size:.75rem; letter-spacing:.07em; text-align:left; text-transform:uppercase; }}
    td {{ padding:1rem; border-top:1px solid var(--line); vertical-align:top; }}
    tbody tr:first-child td {{ border-top:0; }}
    tbody tr:hover {{ background:#faf9ff; }}
    .number {{ width:100px; }}
    .number span, .phase span {{ display:inline-block; border-radius:999px; background:var(--lavender); color:var(--violet-dark); padding:.25rem .55rem; font-size:.78rem; font-weight:800; white-space:nowrap; }}
    .title {{ color:var(--ink); font-size:1.02rem; font-weight:800; text-decoration:none; }}
    .title:hover {{ color:var(--violet-dark); text-decoration:underline; }}
    td p {{ max-width:670px; margin:.28rem 0 0; color:var(--muted); font-size:.9rem; }}
    .phase {{ width:170px; }}
    .phase small {{ display:block; margin:.35rem 0 0 .25rem; color:var(--muted); }}
    .formats {{ width:190px; white-space:nowrap; }}
    .formats a {{ display:inline-block; margin:0 .28rem .35rem 0; border:1px solid #cbbef2; border-radius:6px; padding:.28rem .48rem; font-size:.82rem; font-weight:750; text-decoration:none; }}
    .formats a:hover {{ color:white; background:var(--violet); border-color:var(--violet); }}
    #empty {{ display:none; padding:3rem 1rem; color:var(--muted); text-align:center; }}
    footer {{ width:min(1180px, calc(100% - 2rem)); margin:0 auto 3rem; color:var(--muted); font-size:.87rem; }}
    @media (max-width:760px) {{
      header {{ padding-top:2.6rem; padding-bottom:3.3rem; }} .controls {{ grid-template-columns:1fr; }}
      thead {{ position:absolute; width:1px; height:1px; overflow:hidden; clip:rect(0 0 0 0); }}
      table, tbody, tr, td {{ display:block; width:100% !important; }}
      tr {{ padding:1rem; border-top:1px solid var(--line); }} tr:first-child {{ border-top:0; }} td {{ padding:.3rem 0; border:0 !important; }}
      .formats {{ padding-top:.7rem; }} .summary {{ align-items:flex-start; flex-direction:column; }}
    }}
  </style>
</head>
<body>
  <header>
    <div class="brand"><img src="assets/lightyear-reversed.svg" alt="LIGHTYEAR primary logo"><span>Documentation</span></div>
    <p class="eyebrow">Governed modernization evidence</p>
    <h1>Milestone library</h1>
    <p>Search {len(models)} customer-readable milestone briefs by number, title, customer value, delivered capability, release, or roadmap phase.</p>
  </header>
  <main>
    <section class="controls" aria-label="Milestone filters">
      <div><label for="search">Search milestones</label><input id="search" type="search" placeholder="Try Oracle, stored procedures, CDC, customer pilot…" autocomplete="off"></div>
      <div><label for="phase">Roadmap phase</label><select id="phase"><option value="all">All phases</option><option value="foundation">Foundation (MS #1–10)</option><option value="trust-runtime">Trust and runtime (MS #11–20)</option><option value="pilot-preparation">Pilot preparation (MS #21–32)</option><option value="qualification">Qualification (MS #33–{EXPECTED_MILESTONES[-1]})</option></select></div>
    </section>
    <div class="summary"><span id="result-count" aria-live="polite">{len(models)} milestones</span><a href="{github_blob('docs/milestones/README.md')}">Open repository index</a></div>
    <section class="table-wrap" aria-label="Milestone results">
      <table><thead><tr><th>Milestone</th><th>Customer brief</th><th>Phase / release</th><th>Formats</th></tr></thead><tbody>
        {''.join(rows)}
      </tbody></table>
      <p id="empty">No milestones match those filters. Clear the search or select another phase.</p>
    </section>
  </main>
  <footer>These briefs package committed repository evidence. Underlying receipts, ledgers, gates, tests, and policy decisions remain authoritative.</footer>
  <script>
    (() => {{
      const search = document.querySelector('#search');
      const phase = document.querySelector('#phase');
      const rows = [...document.querySelectorAll('.milestone')];
      const count = document.querySelector('#result-count');
      const empty = document.querySelector('#empty');
      const params = new URLSearchParams(location.search);
      search.value = params.get('q') || '';
      if ([...phase.options].some(option => option.value === params.get('phase'))) phase.value = params.get('phase');
      const apply = () => {{
        const terms = search.value.toLowerCase().trim().split(/\\s+/).filter(Boolean);
        let visible = 0;
        rows.forEach(row => {{
          const matchesText = terms.every(term => row.dataset.search.includes(term));
          const matchesPhase = phase.value === 'all' || row.dataset.phase === phase.value;
          const show = matchesText && matchesPhase;
          row.hidden = !show;
          if (show) visible += 1;
        }});
        count.textContent = `${{visible}} milestone${{visible === 1 ? '' : 's'}}`;
        empty.style.display = visible ? 'none' : 'block';
        const next = new URLSearchParams();
        if (search.value.trim()) next.set('q', search.value.trim());
        if (phase.value !== 'all') next.set('phase', phase.value);
        history.replaceState(null, '', next.size ? `?${{next}}` : location.pathname);
      }};
      search.addEventListener('input', apply);
      phase.addEventListener('change', apply);
      document.addEventListener('keydown', event => {{
        if (event.key === '/' && document.activeElement !== search) {{ event.preventDefault(); search.focus(); }}
        if (event.key === 'Escape' && document.activeElement === search) {{ search.value = ''; apply(); }}
      }});
      apply();
    }})();
  </script>
</body>
</html>
'''
    (DOC_ROOT / "index.html").write_text(page, encoding="utf-8")


def build() -> None:
    catalog, releases = load_catalog(), parse_changelog()
    titles = {entry["number"]: entry["title"] for entry in catalog["milestones"]}
    models = [build_model(entry, releases, titles) for entry in catalog["milestones"]]
    expected = {f"MS-{number:02d}" for number in EXPECTED_MILESTONES}
    for path in DOC_ROOT.glob("MS-*"):
        if path.is_dir() and path.name not in expected:
            shutil.rmtree(path)
    for model in models:
        stem = f"MS-{model['number']:02d}"; directory = DOC_ROOT / stem
        directory.mkdir(parents=True, exist_ok=True)
        (directory / f"{stem}.md").write_text(markdown_text(model, catalog["audience"]), encoding="utf-8")
        build_docx(model, catalog["audience"], directory / f"{stem}.docx")
        build_pdf(model, catalog["audience"], directory / f"{stem}.pdf")
    site_assets = DOC_ROOT / "assets"
    site_assets.mkdir(parents=True, exist_ok=True)
    legacy_site_logo = site_assets / "lightyear-horizontal-reversed.svg"
    if legacy_site_logo.exists():
        legacy_site_logo.unlink()
    shutil.copyfile(BRAND_ASSETS / "lightyear-reversed.svg", site_assets / "lightyear-reversed.svg")
    write_markdown_index(models)
    write_html_index(models)
    artifacts = []
    for number in EXPECTED_MILESTONES:
        stem = f"MS-{number:02d}"
        for suffix in (".md", ".docx", ".pdf"):
            path = DOC_ROOT / stem / f"{stem}{suffix}"
            artifacts.append({"path": path.relative_to(ROOT).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_path(path)})
    library_files = []
    for path in (DOC_ROOT / "README.md", DOC_ROOT / "index.html", DOC_ROOT / "assets" / "lightyear-reversed.svg"):
        library_files.append({"path": path.relative_to(ROOT).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_path(path)})
    manifest = {"schema_version": "1.1", "generator_version": GENERATOR_VERSION, "source_sha256": source_sha256(), "milestone_count": len(EXPECTED_MILESTONES), "artifact_count": EXPECTED_ARTIFACTS, "formats": ["md", "docx", "pdf"], "artifacts": artifacts, "library_files": library_files}
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    print(json.dumps({"status": "built", "milestones": len(EXPECTED_MILESTONES), "artifacts": EXPECTED_ARTIFACTS}, sort_keys=True))


def verify() -> None:
    load_catalog()
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8")); errors = []
    if manifest.get("source_sha256") != source_sha256(): errors.append("documentation sources changed without regeneration")
    if manifest.get("milestone_count") != len(EXPECTED_MILESTONES): errors.append(f"manifest milestone count is not {len(EXPECTED_MILESTONES)}")
    if manifest.get("artifact_count") != EXPECTED_ARTIFACTS: errors.append(f"manifest artifact count is not {EXPECTED_ARTIFACTS}")
    declared = set()
    for artifact in manifest.get("artifacts", []):
        relative = artifact["path"]; declared.add(relative); path = ROOT / relative
        if not path.is_file(): errors.append(f"missing artifact: {relative}"); continue
        if path.stat().st_size != artifact["bytes"]: errors.append(f"size mismatch: {relative}")
        if sha256_path(path) != artifact["sha256"]: errors.append(f"hash mismatch: {relative}")
    actual = {path.relative_to(ROOT).as_posix() for path in DOC_ROOT.glob("MS-*/*") if path.is_file() and path.suffix in {".md", ".docx", ".pdf"}}
    errors.extend(f"undeclared artifact: {path}" for path in sorted(actual - declared))
    errors.extend(f"declared artifact missing: {path}" for path in sorted(declared - actual))
    for library_file in manifest.get("library_files", []):
        relative = library_file["path"]; path = ROOT / relative
        if not path.is_file(): errors.append(f"missing library file: {relative}"); continue
        if path.stat().st_size != library_file["bytes"]: errors.append(f"size mismatch: {relative}")
        if sha256_path(path) != library_file["sha256"]: errors.append(f"hash mismatch: {relative}")
    if errors:
        raise SystemExit("Milestone documentation verification failed:\n- " + "\n- ".join(errors))
    print(json.dumps({"status": "verified", "milestones": len(EXPECTED_MILESTONES), "artifacts": EXPECTED_ARTIFACTS}, sort_keys=True))


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("command", choices=("build", "verify")); args = parser.parse_args()
    build() if args.command == "build" else verify()


if __name__ == "__main__":
    main()
